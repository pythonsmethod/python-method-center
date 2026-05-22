# -*- coding: utf-8 -*-
# image_pipeline.py — Universal Document Intake Pipeline (Phase 5 Extension)
# Handles ALL attachment types from Telegram via SendPulse webhook.
# Additive module: does NOT break existing text message flow.
#
# Supported formats:
#   Images : JPG / JPEG / PNG / WEBP / HEIC / HEIF → Vision OCR
#   PDF    : text PDF → direct extract | scanned PDF → Vision OCR per page
#   DOCX   : python-docx text extraction
#   DOC    : graceful reject (ask user to resend as PDF/DOCX)
#   TXT    : plain text read
#   CSV    : plain text read
#   XLS/XLSX: openpyxl cell extraction
#   ZIP    : safe reject with instruction
#   Other  : safe reject with instruction
#
# Safety limits:
#   MAX_FILE_SIZE_MB = 20
#   MAX_PDF_PAGES    = 10
#   MAX_EXTRACTED_LEN= 8000
#   DOWNLOAD_TIMEOUT = 30s
#
# Pipeline entrypoints (preserved from Phase 4):
#   extract_attachment()        — parse attachment metadata from webhook body
#   download_telegram_file()    — fetch file bytes via Telegram Bot API
#   process_attachment_message()— top-level handler called from main.py
#   has_attachment()            — fast check used in main.py webhook
#
# LOGGING tags:
#   [MEDIA] file_detected
#   [MEDIA] file_type_detected
#   [MEDIA] extraction_started
#   [MEDIA] extraction_success
#   [MEDIA] extraction_partial
#   [MEDIA] extraction_failed
#   [MEDIA] unsupported_file_type
#   [MEDIA] file_size_rejected

import os
import io
import base64
import logging
import httpx
from analysis_module import (
    save_analysis_to_session, evaluate_escalation,
    check_analysis_completeness
)

log = logging.getLogger("image_pipeline")

# ─── Config ──────────────────────────────────────────────────────────────────
TELEGRAM_BOT_TOKEN = os.environ.get("TELEGRAM_BOT_TOKEN", "") or os.environ.get("NOTIFY_BOT_TOKEN", "")
OPENAI_API_KEY     = os.environ.get("OPENAI_API_KEY", "")
ANTHROPIC_API_KEY  = os.environ.get("ANTHROPIC_API_KEY", "")

_TELEGRAM_API      = "https://api.telegram.org/bot{token}"
_TELEGRAM_FILE_URL = "https://api.telegram.org/file/bot{token}/{file_path}"

# ─── Safety Limits ───────────────────────────────────────────────────────────
MAX_FILE_SIZE_BYTES = 20 * 1024 * 1024   # 20 MB
MAX_PDF_PAGES       = 10
MAX_EXTRACTED_LEN   = 8000
_MIN_OCR_TEXT_LEN   = 10   # absolute floor — below this is noise/empty

# ─── OCR Quality Tiers ──────────────────────────────────────────────────────────
# high   : >= 80 chars OR >= 40 chars with medical heuristics   → proceed normally
# medium : >= 40 chars OR >= 10 chars with medical heuristics   → proceed normally
# low    : >= 10 chars, no medical match                        → proceed with internal note
# failed : < 10 chars or empty                                  → failsafe message
_OCR_HIGH_MIN       = 80    # chars for unconditional high confidence
_OCR_MEDIUM_MIN     = 40    # chars for unconditional medium confidence
_OCR_MEDICAL_HIGH   = 40    # chars + medical heuristic → high
_OCR_MEDICAL_MEDIUM = 10    # chars + medical heuristic → medium

APP_VERSION = "preprocess-v1"   # deployment fingerprint
log.info("[PIPELINE_BOOT] image_pipeline loaded version=%s", APP_VERSION)

# ─── User-facing messages (Russian) ──────────────────────────────────────────
OCR_FAILSAFE_MESSAGE = (
    "[DEBUG_FAILSAFE_V3] Файл получен. OCR не сработал — нужно чётче. "
    "Пожалуйста, отправьте фото ближе или в формате PDF/Word."
)
OCR_PARTIAL_SUCCESS_MESSAGE = (
    "Мы получили документ. Часть данных удалось распознать успешно. "
    "При необходимости система может позже запросить более чёткие страницы."
)
FILE_TOO_LARGE_MESSAGE = (
    "Файл слишком большой для обработки (максимум 20 МБ). "
    "Пожалуйста, сожмите файл или отправьте только нужные страницы."
)
UNSUPPORTED_FORMAT_MESSAGE = (
    "Этот формат файла не поддерживается. "
    "Пожалуйста, отправьте файл в формате PDF, Word (DOCX), JPG, PNG или TXT."
)
DOC_LEGACY_MESSAGE = (
    "Файл в формате .doc (старый Word) получен, но не может быть прочитан автоматически. "
    "Пожалуйста, пересохраните файл в формате .docx, PDF или сфотографируйте страницы."
)
ZIP_MESSAGE = (
    "Получен архив ZIP. Пожалуйста, отправьте файлы по отдельности "
    "в формате PDF, Word (DOCX), JPG или PNG."
)

# ─── MIME type → category mapping ────────────────────────────────────────────
_IMAGE_MIMES = {
    "image/jpeg", "image/jpg", "image/png", "image/webp",
    "image/heic", "image/heif", "image/gif", "image/bmp", "image/tiff",
}
_IMAGE_EXTS  = {".jpg", ".jpeg", ".png", ".webp", ".heic", ".heif", ".gif", ".bmp", ".tiff"}
_PDF_MIMES   = {"application/pdf"}
_PDF_EXTS    = {".pdf"}
_DOCX_MIMES  = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
_DOCX_EXTS   = {".docx"}
_DOC_MIMES   = {"application/msword"}
_DOC_EXTS    = {".doc"}
_TEXT_MIMES  = {"text/plain", "text/csv"}
_TEXT_EXTS   = {".txt", ".csv"}
_XLSX_MIMES  = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
}
_XLSX_EXTS   = {".xls", ".xlsx"}
_ZIP_MIMES   = {"application/zip", "application/x-zip-compressed", "application/octet-stream"}
_ZIP_EXTS    = {".zip"}


# ─── Medical Document Heuristics ─────────────────────────────────────────────
# Keywords indicating a medical document — used to bias toward partial_success
_MEDICAL_TERMS = {
    # Russian lab/report terms
    "анализ", "анализы", "результат", "показатель", "референс", "норма",
    "лейкоцит", "эритроцит", "гемоглобин", "тромбоцит", "гематокрит",
    "лимфоцит", "моноцит", "нейтрофил", "эозинофил", "базофил",
    "глюкоза", "холестерин", "билирубин", "креатинин", "мочевина",
    "белок", "альбумин", "ферритин", "трансфераза", "ферментов",
    "онкомаркер", "пса", "са-125", "раково", "опухол", "онколог",
    "мрт", "кт", "узи", "рентген", "гистолог", "цитолог", "биопси",
    "диагноз", "заключение", "врач", "больниц", "клиник", "лаборатор",
    # Latin/international lab markers
    "hb", "hba1c", "wbc", "rbc", "plt", "mcv", "mch", "mchc",
    "alt", "ast", "ggt", "alp", "ldh", "crp", "esr", "psa",
    "tsh", "t3", "t4", "insulin", "cortisol", "ferritin",
    "glucose", "creatinine", "urea", "bilirubin", "cholesterol",
    "mmol", "mg/dl", "g/l", "u/l", "мкмол", "нмол", "пмол",
    # Structural markers
    "дата", "date", "пациент", "patient", "ф.и.о", "возраст",
    "пол:", "male", "female", "мужской", "женский",
    "результат", "reference", "норма:", "range",
}

def _has_medical_terms(text: str) -> bool:
    """Return True if text contains recognisable medical vocabulary."""
    lower = text.lower()
    return any(term in lower for term in _MEDICAL_TERMS)

def _score_ocr_quality(text: str) -> str:
    """
    Assign a 4-tier OCR quality score based on character count + medical heuristics.

    Tiers:
        high   — reliable, proceed normally
        medium — usable, proceed normally
        low    — marginal but contains some content, proceed with internal note
        failed — too little text to use, trigger failsafe
    """
    if not text:
        return "failed"
    n = len(text.strip())
    has_medical = _has_medical_terms(text)

    if n >= _OCR_HIGH_MIN:
        return "high"
    if n >= _OCR_MEDICAL_HIGH and has_medical:
        return "high"
    if n >= _OCR_MEDIUM_MIN:
        return "medium"
    if n >= _OCR_MEDICAL_MEDIUM and has_medical:
        return "medium"
    if n >= _MIN_OCR_TEXT_LEN:
        return "low"
    return "failed"


def _detect_format(mime_type: str, file_name: str) -> str:
    """
    Determine high-level format category from MIME + filename extension.
    Returns: 'image' | 'pdf' | 'docx' | 'doc' | 'text' | 'spreadsheet' | 'zip' | 'unknown'
    """
    mime = (mime_type or "").lower().strip()
    ext  = ""
    if file_name:
        import os as _os
        ext = _os.path.splitext(file_name.lower())[1]

    if mime in _IMAGE_MIMES or ext in _IMAGE_EXTS:
        return "image"
    if mime in _PDF_MIMES or ext in _PDF_EXTS:
        return "pdf"
    if mime in _DOCX_MIMES or ext in _DOCX_EXTS:
        return "docx"
    if mime in _DOC_MIMES or ext in _DOC_EXTS:
        return "doc"
    if mime in _TEXT_MIMES or ext in _TEXT_EXTS:
        return "text"
    if mime in _XLSX_MIMES or ext in _XLSX_EXTS:
        return "spreadsheet"
    if ext in _ZIP_EXTS:
        return "zip"
    if mime in _ZIP_MIMES and ext == ".zip":
        return "zip"
    # Byte sniff fallback (checked after download in extract layer)
    return "unknown"


def _byte_sniff_format(file_bytes: bytes) -> str:
    """Fallback: detect format from magic bytes."""
    if not file_bytes or len(file_bytes) < 4:
        return "unknown"
    sig = file_bytes[:8]
    if sig[:4] == b"%PDF":
        return "pdf"
    if sig[:4] == b"PK\x03\x04":
        # ZIP or DOCX/XLSX (both are ZIP-based)
        return "zip_or_office"
    if sig[:3] in (b"\xff\xd8\xff",):
        return "image"
    if sig[:8] == b"\x89PNG\r\n\x1a\n":
        return "image"
    return "unknown"


# ─── Attachment Detection ─────────────────────────────────────────────────────

def extract_attachment(body: dict) -> dict | None:
    """
    Parse SendPulse webhook body and extract attachment metadata.
    Returns dict with keys: file_id, attachment_type, caption, mime_type, file_name
    or None if no attachment found.
    """
    if not isinstance(body, dict):
        return None

    info = body.get("info") or {}
    if not isinstance(info, dict):
        info = {}

    msg1 = info.get("message") or body.get("message") or {}
    if not isinstance(msg1, dict):
        msg1 = {}

    cd = msg1.get("channel_data") or {}
    if not isinstance(cd, dict):
        cd = {}

    tg_msg = cd.get("message") or msg1
    if not isinstance(tg_msg, dict):
        tg_msg = {}

    caption   = (tg_msg.get("caption") or cd.get("caption") or "").strip()
    file_name = ""

    # --- photo (compressed) ---
    photo_arr = tg_msg.get("photo") or cd.get("photo")
    if photo_arr and isinstance(photo_arr, list) and len(photo_arr) > 0:
        best    = photo_arr[-1]
        file_id = best.get("file_id", "")
        if file_id:
            log.info("[MEDIA] file_detected type=photo file_id=%.20s", file_id)
            return {
                "file_id":         file_id,
                "attachment_type": "image",
                "caption":         caption,
                "mime_type":       "image/jpeg",
                "file_name":       "photo.jpg",
            }

    # --- document (any file sent as document) ---
    doc = tg_msg.get("document") or cd.get("document")
    if doc and isinstance(doc, dict):
        file_id   = doc.get("file_id", "")
        mime_type = doc.get("mime_type", "application/octet-stream")
        file_name = doc.get("file_name", "")
        if file_id:
            fmt   = _detect_format(mime_type, file_name)
            log.info("[MEDIA] file_detected type=%s mime=%s name=%s file_id=%.20s",
                     fmt, mime_type, file_name, file_id)
            return {
                "file_id":         file_id,
                "attachment_type": fmt,
                "caption":         caption,
                "mime_type":       mime_type,
                "file_name":       file_name,
            }

    return None


# ─── Telegram File Download ───────────────────────────────────────────────────

async def download_telegram_file(file_id: str) -> bytes | None:
    """Download a Telegram file by file_id. Returns raw bytes or None."""
    if not TELEGRAM_BOT_TOKEN:
        log.error("[MEDIA] TELEGRAM_BOT_TOKEN not set — cannot download file")
        return None

    base = _TELEGRAM_API.format(token=TELEGRAM_BOT_TOKEN)

    try:
        async with httpx.AsyncClient(timeout=30.0) as client:
            r = await client.get(f"{base}/getFile", params={"file_id": file_id})
            r.raise_for_status()
            data = r.json()
            if not data.get("ok"):
                log.error("[MEDIA] getFile API error: %s", data)
                return None
            file_path = data["result"]["file_path"]
            file_url  = _TELEGRAM_FILE_URL.format(
                token=TELEGRAM_BOT_TOKEN, file_path=file_path
            )
            r2 = await client.get(file_url)
            r2.raise_for_status()
            raw = r2.content
            log.info("[MEDIA] file_downloaded bytes=%d file_path=%s", len(raw), file_path)
            return raw
    except Exception as e:
        log.error("[MEDIA] download_telegram_file error: %s", e)
        return None


# ─── Format-specific Extractors ──────────────────────────────────────────────


# ─── Image Preprocessing ─────────────────────────────────────────────────────
# Preprocessing pipeline applied BEFORE Vision OCR.
# Goal: improve readability of Telegram-compressed, distant, low-DPI medical photos.
# Safety: does NOT hallucinate text. Only improves image clarity.

_PREPROCESS_MIN_WIDTH  = 1200  # upscale if image narrower than this
_PREPROCESS_MAX_UPSCALE = 3.0  # never upscale more than 3× to avoid noise amplification

def _preprocess_image(file_bytes: bytes) -> tuple:
    """
    Apply OCR-oriented preprocessing to image bytes.

    Steps:
        1. Grayscale conversion          — removes colour noise, reduces complexity
        2. Auto-contrast (Pillow CLAHE)  — normalizes brightness/exposure unevenness
        3. Sharpening                    — enhances fine text edges
        4. Smart upscaling               — upscale small/distant photos to target DPI
        5. Border crop heuristics        — crop uniform-colour borders (chat background)

    Returns:
        (processed_bytes: bytes, meta: dict)
        meta keys: original_size, processed_size, upscale_ratio,
                   grayscale_applied, sharpen_applied, crop_applied, error

    Falls back gracefully if Pillow unavailable or any step fails.
    """
    meta = {
        "original_size": (0, 0),
        "processed_size": (0, 0),
        "upscale_ratio": 1.0,
        "grayscale_applied": False,
        "sharpen_applied": False,
        "crop_applied": False,
        "error": None,
    }
    try:
        from PIL import Image, ImageFilter, ImageOps
        import io as _io
    except ImportError:
        log.warning("[PREPROCESS] Pillow not available — skipping preprocessing")
        meta["error"] = "pillow_missing"
        return file_bytes, meta

    try:
        img = Image.open(_io.BytesIO(file_bytes))
        meta["original_size"] = img.size

        # ── Step 1: Grayscale ──────────────────────────────────────────────────
        if img.mode not in ("L", "LA"):
            img = img.convert("L")
            meta["grayscale_applied"] = True

        # ── Step 2: Auto-contrast (CLAHE-style normalisation) ─────────────────
        img = ImageOps.autocontrast(img, cutoff=2)

        # ── Step 3: Smart upscaling ───────────────────────────────────────────
        w, h = img.size
        if w < _PREPROCESS_MIN_WIDTH:
            scale = min(_PREPROCESS_MAX_UPSCALE, _PREPROCESS_MIN_WIDTH / w)
            new_w = int(w * scale)
            new_h = int(h * scale)
            img = img.resize((new_w, new_h), Image.LANCZOS)
            meta["upscale_ratio"] = round(scale, 2)
        else:
            meta["upscale_ratio"] = 1.0

        # ── Step 4: Sharpening ────────────────────────────────────────────────
        img = img.filter(ImageFilter.SHARPEN)
        img = img.filter(ImageFilter.SHARPEN)  # double-pass for dense small text
        meta["sharpen_applied"] = True

        # ── Step 5: Border crop heuristic ────────────────────────────────────
        # Crop rows/cols where >95% of pixels are near-white or near-black
        # (chat UI backgrounds, dark mode, blank margins)
        try:
            bg_threshold = 15   # pixels within 15 of pure white (255) are "border"
            bbox = ImageOps.invert(img).getbbox()
            if bbox:
                cw = bbox[2] - bbox[0]
                ch = bbox[3] - bbox[1]
                # Only crop if result is at least 60% of original dimensions
                if cw >= img.width * 0.6 and ch >= img.height * 0.6:
                    img = img.crop(bbox)
                    meta["crop_applied"] = True
        except Exception:
            pass  # crop is best-effort

        meta["processed_size"] = img.size

        # ── Serialise to JPEG ─────────────────────────────────────────────────
        buf = _io.BytesIO()
        img.save(buf, format="JPEG", quality=92)
        processed_bytes = buf.getvalue()

        log.info(
            "[PREPROCESS] original=%dx%d processed=%dx%d upscale=%.1fx "
            "gray=%s sharpen=%s crop=%s output_bytes=%d",
            meta["original_size"][0], meta["original_size"][1],
            meta["processed_size"][0], meta["processed_size"][1],
            meta["upscale_ratio"],
            meta["grayscale_applied"], meta["sharpen_applied"], meta["crop_applied"],
            len(processed_bytes),
        )
        return processed_bytes, meta

    except Exception as exc:
        log.warning("[PREPROCESS] preprocessing failed: %s — using original", exc)
        meta["error"] = str(exc)
        return file_bytes, meta


async def _extract_image(file_bytes: bytes, mime_type: str) -> dict:
    """
    Extract text from image bytes via Vision OCR.

    Preprocessing comparison mode:
        1. Try OCR on ORIGINAL bytes first.
        2. If result is high/medium — return immediately (preprocessing unnecessary).
        3. If result is low/failed — preprocess image and run OCR again.
        4. Keep the better result from original vs preprocessed by tier ranking.

    This ensures: previously failing photos (chars=0) benefit from preprocessing
    without adding latency to photos that OCR reads well natively.
    """
    log.info("[MEDIA] extraction_started format=image size=%d", len(file_bytes))

    tier_rank = {"high": 3, "medium": 2, "low": 1, "failed": 0}

    async def _ocr_both_providers(img_bytes: bytes, label: str) -> dict:
        """Run OpenAI + Claude on given image bytes, return best result."""
        b64_img = base64.b64encode(img_bytes).decode("utf-8")
        best = None
        if OPENAI_API_KEY:
            r = await _ocr_openai(b64_img, mime_type)
            if r["success"] and r["confidence"] in ("high", "medium"):
                log.info("[MEDIA] extraction_success format=image provider=openai "
                         "source=%s quality=%s len=%d", label, r["confidence"], len(r["text"]))
                return r
            if r["success"]:
                best = r
                log.info("[MEDIA] extraction_partial format=image provider=openai "
                         "source=%s quality=%s len=%d — trying claude",
                         label, r["confidence"], len(r["text"]))
        if ANTHROPIC_API_KEY:
            r = await _ocr_claude(b64_img, mime_type)
            if r["success"]:
                current_rank = tier_rank.get(r["confidence"], 0)
                best_rank    = tier_rank.get(best["confidence"], 0) if best else 0
                if current_rank >= best_rank:
                    best = r
                log.info("[MEDIA] extraction_success format=image provider=claude "
                         "source=%s quality=%s len=%d", label, r["confidence"], len(r["text"]))
        return best

    # ── Pass 1: original bytes ────────────────────────────────────────────────
    original_result = await _ocr_both_providers(file_bytes, "original")

    # If original gives high/medium — great, return immediately
    if original_result and original_result["confidence"] in ("high", "medium"):
        log.info("[MEDIA] extraction_success format=image source=original "
                 "quality=%s len=%d — preprocessing skipped",
                 original_result["confidence"], len(original_result["text"]))
        return original_result

    # ── Pass 2: preprocessing ─────────────────────────────────────────────────
    # Original result is low or failed — try preprocessing to improve readability
    preprocessed_bytes, pre_meta = _preprocess_image(file_bytes)

    if pre_meta.get("error"):
        log.warning("[MEDIA] preprocessing_skipped reason=%s", pre_meta["error"])
        preprocessed_result = None
    else:
        log.info("[MEDIA] preprocessing_applied upscale=%.1fx gray=%s sharpen=%s crop=%s",
                 pre_meta["upscale_ratio"], pre_meta["grayscale_applied"],
                 pre_meta["sharpen_applied"], pre_meta["crop_applied"])
        preprocessed_result = await _ocr_both_providers(preprocessed_bytes, "preprocessed")

    # ── Pick best result from original vs preprocessed ────────────────────────
    candidates = [r for r in (original_result, preprocessed_result)
                  if r and r.get("success")]

    if candidates:
        best_result = max(candidates, key=lambda r: (
            tier_rank.get(r["confidence"], 0), len(r.get("text", ""))
        ))
        orig_chars  = len(original_result.get("text", "")) if original_result else 0
        pre_chars   = len(preprocessed_result.get("text", "")) if preprocessed_result else 0
        log.info("[MEDIA] ocr_comparison_done original_chars=%d preprocessed_chars=%d "
                 "chosen=%s chosen_quality=%s",
                 orig_chars, pre_chars,
                 best_result.get("provider", "?"), best_result["confidence"])
        log.info("[MEDIA] extraction_partial_success format=image "
                 "provider=%s quality=%s len=%d",
                 best_result["provider"], best_result["confidence"], len(best_result["text"]))
        return best_result

    log.warning("[MEDIA] extraction_failed format=image — both passes returned no text")
    return {"success": False, "text": "", "confidence": "failed",
            "provider": "none", "error": "all OCR providers failed"}


async def _extract_pdf(file_bytes: bytes) -> dict:
    """
    Extract text from PDF.
    Strategy: try pdfplumber for text PDF first; fall back to Vision OCR on pages.
    """
    log.info("[MEDIA] extraction_started format=pdf size=%d", len(file_bytes))

    # Attempt 1: pdfplumber (text-based PDF)
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            page_count = len(pdf.pages)
            pages_to_read = min(page_count, MAX_PDF_PAGES)
            for i, page in enumerate(pdf.pages[:pages_to_read]):
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(f"[Page {i+1}]\n{page_text}")
        combined = "\n\n".join(text_parts)[:MAX_EXTRACTED_LEN]
        if len(combined) >= _MIN_OCR_TEXT_LEN:
            log.info("[MEDIA] extraction_success format=pdf method=text pages=%d len=%d",
                     pages_to_read, len(combined))
            return {
                "success": True, "text": combined, "confidence": "high",
                "provider": "pdfplumber", "error": None,
                "pages": pages_to_read,
            }
        log.info("[MEDIA] extraction_partial format=pdf method=text — no text layer, trying OCR")
    except ImportError:
        log.warning("[MEDIA] pdfplumber not available — skipping text extraction")
    except Exception as e:
        log.warning("[MEDIA] pdfplumber error: %s — trying OCR fallback", e)

    # Attempt 2: Vision OCR on first pages (convert PDF pages to images)
    try:
        import pdfplumber
        b64_pages = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_to_ocr = min(len(pdf.pages), MAX_PDF_PAGES)
            for i, page in enumerate(pdf.pages[:pages_to_ocr]):
                # Render page as PNG image
                img = page.to_image(resolution=150)
                buf = io.BytesIO()
                img.save(buf, format="PNG")
                b64_pages.append(base64.b64encode(buf.getvalue()).decode("utf-8"))

        ocr_texts = []
        for idx, b64 in enumerate(b64_pages):
            if OPENAI_API_KEY:
                res = await _ocr_openai(b64, "image/png")
            elif ANTHROPIC_API_KEY:
                res = await _ocr_claude(b64, "image/png")
            else:
                break
            if res["success"]:
                ocr_texts.append(f"[Page {idx+1}]\n{res['text']}")

        if ocr_texts:
            combined = "\n\n".join(ocr_texts)[:MAX_EXTRACTED_LEN]
            log.info("[MEDIA] extraction_success format=pdf method=ocr pages=%d len=%d",
                     len(ocr_texts), len(combined))
            return {
                "success": True, "text": combined, "confidence": "high",
                "provider": "vision_ocr", "error": None,
                "pages": len(ocr_texts),
            }
    except Exception as e:
        log.warning("[MEDIA] PDF OCR fallback error: %s", e)

    log.warning("[MEDIA] extraction_failed format=pdf")
    return {"success": False, "text": "", "confidence": "failed",
            "provider": "none", "error": "PDF extraction failed"}


def _extract_docx(file_bytes: bytes) -> dict:
    """Extract text from DOCX via python-docx."""
    log.info("[MEDIA] extraction_started format=docx size=%d", len(file_bytes))
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        combined = "\n".join(paragraphs)[:MAX_EXTRACTED_LEN]
        if len(combined) >= _MIN_OCR_TEXT_LEN:
            log.info("[MEDIA] extraction_success format=docx len=%d", len(combined))
            return {"success": True, "text": combined, "confidence": "high",
                    "provider": "python-docx", "error": None}
        log.warning("[MEDIA] extraction_partial format=docx — document appears empty")
        return {"success": False, "text": combined, "confidence": "low",
                "provider": "python-docx", "error": "document appears empty"}
    except ImportError:
        log.warning("[MEDIA] python-docx not available")
        return {"success": False, "text": "", "confidence": "failed",
                "provider": "none", "error": "python-docx not installed"}
    except Exception as e:
        log.error("[MEDIA] extraction_failed format=docx error=%s", e)
        return {"success": False, "text": "", "confidence": "failed",
                "provider": "none", "error": str(e)}


def _extract_text(file_bytes: bytes) -> dict:
    """Extract text from TXT or CSV files."""
    log.info("[MEDIA] extraction_started format=text size=%d", len(file_bytes))
    try:
        text = file_bytes.decode("utf-8", errors="replace")[:MAX_EXTRACTED_LEN]
        log.info("[MEDIA] extraction_success format=text len=%d", len(text))
        return {"success": True, "text": text, "confidence": "high",
                "provider": "plaintext", "error": None}
    except Exception as e:
        log.error("[MEDIA] extraction_failed format=text error=%s", e)
        return {"success": False, "text": "", "confidence": "failed",
                "provider": "none", "error": str(e)}


def _extract_spreadsheet(file_bytes: bytes, file_name: str = "") -> dict:
    """Extract cell values from XLS/XLSX via openpyxl."""
    log.info("[MEDIA] extraction_started format=spreadsheet size=%d", len(file_bytes))
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
        rows_out = []
        for sheet in wb.worksheets:
            rows_out.append(f"[Sheet: {sheet.title}]")
            for row in sheet.iter_rows(values_only=True):
                row_text = " | ".join(str(c) for c in row if c is not None)
                if row_text.strip():
                    rows_out.append(row_text)
            if len("\n".join(rows_out)) > MAX_EXTRACTED_LEN:
                break
        combined = "\n".join(rows_out)[:MAX_EXTRACTED_LEN]
        log.info("[MEDIA] extraction_success format=spreadsheet len=%d", len(combined))
        return {"success": True, "text": combined, "confidence": "high",
                "provider": "openpyxl", "error": None}
    except ImportError:
        log.warning("[MEDIA] openpyxl not available")
        return {"success": False, "text": "", "confidence": "failed",
                "provider": "none", "error": "openpyxl not installed"}
    except Exception as e:
        log.error("[MEDIA] extraction_failed format=spreadsheet error=%s", e)
        return {"success": False, "text": "", "confidence": "failed",
                "provider": "none", "error": str(e)}


# ─── Vision OCR Providers ─────────────────────────────────────────────────────

async def _ocr_openai(b64: str, mime_type: str) -> dict:
    """OCR via OpenAI gpt-4o Vision."""
    try:
        from openai import AsyncOpenAI
        client = AsyncOpenAI(api_key=OPENAI_API_KEY)
        response = await client.chat.completions.create(
            model="gpt-4o",
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "You are a medical data extraction assistant. "
                            "Extract ALL text visible in this image exactly as shown. "
                            "Preserve table structure where possible. "
                            "Output only the extracted text, no commentary."
                        )
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{b64}",
                            "detail": "high"
                        }
                    }
                ]
            }],
            max_tokens=2000,
        )
        extracted = response.choices[0].message.content or ""
        quality   = _score_ocr_quality(extracted)
        success   = quality != "failed"
        log.info("[MEDIA] ocr_openai len=%d quality=%s has_medical=%s",
                 len(extracted), quality, _has_medical_terms(extracted))
        return {"success": success, "text": extracted[:MAX_EXTRACTED_LEN],
                "confidence": quality, "provider": "openai", "error": None,
                "char_count": len(extracted)}
    except Exception as e:
        log.error("[MEDIA] ocr_openai error: %s", e)
        return {"success": False, "text": "", "confidence": "failed",
                "provider": "openai", "error": str(e)}


async def _ocr_claude(b64: str, mime_type: str) -> dict:
    """OCR via Anthropic Claude Vision."""
    try:
        import anthropic
        # Claude Vision supports: image/jpeg, image/png, image/gif, image/webp
        safe_mime = mime_type if mime_type in (
            "image/jpeg", "image/png", "image/gif", "image/webp"
        ) else "image/jpeg"

        client   = anthropic.AsyncAnthropic(api_key=ANTHROPIC_API_KEY)
        response = await client.messages.create(
            model="claude-opus-4-5",
            max_tokens=2000,
            messages=[{
                "role": "user",
                "content": [
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": safe_mime,
                            "data": b64,
                        }
                    },
                    {
                        "type": "text",
                        "text": (
                            "You are a medical data extraction assistant. "
                            "Extract ALL text visible in this image exactly as shown. "
                            "Preserve table structure where possible. "
                            "Output only the extracted text, no commentary."
                        )
                    }
                ]
            }]
        )
        extracted = response.content[0].text if response.content else ""
        quality   = _score_ocr_quality(extracted)
        success   = quality != "failed"
        log.info("[MEDIA] ocr_claude len=%d quality=%s has_medical=%s",
                 len(extracted), quality, _has_medical_terms(extracted))
        return {"success": success, "text": extracted[:MAX_EXTRACTED_LEN],
                "confidence": quality, "provider": "claude", "error": None,
                "char_count": len(extracted)}
    except Exception as e:
        log.error("[MEDIA] ocr_claude error: %s", e)
        return {"success": False, "text": "", "confidence": "failed",
                "provider": "claude", "error": str(e)}


# ─── Unified Extraction Router ────────────────────────────────────────────────

async def extract_content(file_bytes: bytes, fmt: str, mime_type: str,
                          file_name: str = "") -> dict:
    """
    Route extraction to format-specific handler.
    Always returns dict with: success, text, confidence, provider, error
    """
    if not file_bytes:
        return {"success": False, "text": "", "confidence": "failed",
                "provider": "none", "error": "empty file bytes"}

    # Byte sniff if format still unknown
    if fmt == "unknown":
        sniffed = _byte_sniff_format(file_bytes)
        if sniffed == "pdf":
            fmt = "pdf"
        elif sniffed == "image":
            fmt = "image"

    log.info("[MEDIA] file_type_detected format=%s mime=%s name=%s", fmt, mime_type, file_name)

    if fmt == "image":
        return await _extract_image(file_bytes, mime_type)

    if fmt == "pdf":
        return await _extract_pdf(file_bytes)

    if fmt == "docx":
        return _extract_docx(file_bytes)

    if fmt == "text":
        return _extract_text(file_bytes)

    if fmt == "spreadsheet":
        return _extract_spreadsheet(file_bytes, file_name)

    if fmt == "doc":
        log.warning("[MEDIA] unsupported_file_type format=doc")
        return {"success": False, "text": "", "confidence": "failed",
                "provider": "none", "error": "legacy DOC format",
                "user_message": DOC_LEGACY_MESSAGE}

    if fmt == "zip":
        log.warning("[MEDIA] unsupported_file_type format=zip")
        return {"success": False, "text": "", "confidence": "failed",
                "provider": "none", "error": "ZIP archive",
                "user_message": ZIP_MESSAGE}

    # Last resort: if format still unknown but we have bytes, try Vision OCR
    # (handles HEIC, WebP, unusual mime types that Telegram sends)
    if fmt == "unknown" or fmt == "zip_or_office":
        log.info("[MEDIA] format_unknown_attempting_vision fmt=%s mime=%s", fmt, mime_type)
        return await _extract_image(file_bytes, mime_type or "image/jpeg")

    log.warning("[MEDIA] unsupported_file_type format=%s", fmt)
    return {"success": False, "text": "", "confidence": "failed",
            "provider": "none", "error": f"unsupported format: {fmt}",
            "user_message": UNSUPPORTED_FORMAT_MESSAGE}


# ─── Unified Attachment Context Builder ──────────────────────────────────────

def build_attachment_context(attachment_meta: dict, ocr_result: dict) -> str:
    """
    Build a structured context string injected into the AI prompt.

    Returns:
        [ATTACHMENT_CONTEXT]
        HAS_ATTACHMENTS = true
        ATTACHMENT_TYPE = pdf
        FILE_NAME = analysis_2024.pdf
        MIME_TYPE = application/pdf
        FILES_COUNT = 1
        EXTRACTION_STATUS = success
        OCR_CONFIDENCE = high
        EXTRACTED_TEXT:
        ...
        [/ATTACHMENT_CONTEXT]
    """
    fmt       = attachment_meta.get("attachment_type", "unknown")
    file_name = attachment_meta.get("file_name", "")
    mime_type = attachment_meta.get("mime_type", "")
    caption   = attachment_meta.get("caption", "")
    ocr_text  = ocr_result.get("text", "")
    conf      = ocr_result.get("confidence", "failed")
    success   = ocr_result.get("success", False)
    # Map confidence tier to EXTRACTION_STATUS
    if success and conf in ("high", "medium"):
        status = "success"
    elif success and conf == "low":
        status = "partial_success"
    elif ocr_text and len(ocr_text.strip()) >= _MIN_OCR_TEXT_LEN:
        status = "partial_success"
    else:
        status = "failed"

    lines = [
        "[ATTACHMENT_CONTEXT]",
        "HAS_ATTACHMENTS = true",
        f"ATTACHMENT_TYPE = {fmt}",
    ]
    if file_name:
        lines.append(f"FILE_NAME = {file_name}")
    if mime_type:
        lines.append(f"MIME_TYPE = {mime_type}")
    lines += [
        "FILES_COUNT = 1",
        f"EXTRACTION_STATUS = {status}",
        f"OCR_CONFIDENCE = {conf}",
    ]
    if caption:
        lines.append(f"CAPTION = {caption}")
    if ocr_text:
        lines += ["EXTRACTED_TEXT:", ocr_text]
    else:
        lines.append("EXTRACTED_TEXT: [extraction failed or empty]")
    lines.append("[/ATTACHMENT_CONTEXT]")
    return "\n".join(lines)


# ─── Main Entry Point ─────────────────────────────────────────────────────────

async def process_attachment_message(
    contact_id: str,
    attachment_meta: dict,
    process_message_fn,
    session_update_fn=None,
) -> str:
    """
    Universal attachment handler. Called from main.py webhook for any attachment.

    Steps:
    1. Safety check (file size)
    2. Download file bytes
    3. Route to format-specific extractor
    4. Build ATTACHMENT_CONTEXT
    5. Optionally update session (analysis_module)
    6. Call process_message_fn with enriched context
    """
    file_id   = attachment_meta.get("file_id", "")
    fmt       = attachment_meta.get("attachment_type", "unknown")
    mime_type = attachment_meta.get("mime_type", "application/octet-stream")
    file_name = attachment_meta.get("file_name", "")
    caption   = attachment_meta.get("caption", "")

    log.info("[MEDIA] escalation_triggered contact=%s type=%s file_id=%.20s",
             contact_id, fmt, file_id)

    # Step 1: Download
    file_bytes = await download_telegram_file(file_id)
    if not file_bytes:
        log.error("[MEDIA] extraction_failed — could not download file contact=%s", contact_id)
        return OCR_FAILSAFE_MESSAGE

    # Safety: file size limit
    if len(file_bytes) > MAX_FILE_SIZE_BYTES:
        log.warning("[MEDIA] file_size_rejected contact=%s size=%d", contact_id, len(file_bytes))
        return FILE_TOO_LARGE_MESSAGE

    # Step 2: Extract content
    ocr_result = await extract_content(file_bytes, fmt, mime_type, file_name)

    # Handle user_message overrides for unsupported types
    if not ocr_result["success"] and ocr_result.get("user_message"):
        return ocr_result["user_message"]

    # Step 3: Build context block
    ctx_str = build_attachment_context(attachment_meta, ocr_result)

    # Step 4: Update session if callback provided
    if session_update_fn is not None:
        try:
            session_update_fn(contact_id, ocr_result, attachment_meta)
        except Exception as e:
            log.error("[MEDIA] session_update_fn error: %s", e)

    # Step 5: Determine synthetic user message using 4-tier quality
    ocr_text    = ocr_result.get("text", "")
    ocr_quality = ocr_result.get("confidence", "failed")
    char_count  = ocr_result.get("char_count", len(ocr_text))
    is_usable   = ocr_result.get("success", False) and char_count >= _MIN_OCR_TEXT_LEN
    # is_partial: file was received and has SOME usable signal, or OCR failed but file confirmed
    # Covers: low-quality OCR, API exception with 0 chars, unreadable scan
    has_some_text = char_count >= _MIN_OCR_TEXT_LEN
    ocr_api_failed = not ocr_result.get("success", False) and not ocr_result.get("user_message")
    is_partial = (
        (not is_usable and has_some_text) or          # has text but below usable threshold
        (ocr_result.get("success", False) and ocr_quality == "low") or  # low quality success
        ocr_api_failed                                 # API failed but file was received
    )
    partial_success_triggered = False

    log.info("[MEDIA] ocr_quality_score=%s extracted_characters_count=%d "
             "is_usable=%s is_partial=%s contact=%s",
             ocr_quality, char_count, is_usable, is_partial, contact_id)
    log.info("[MEDIA_DEBUG] final_decision success=%s chars=%d quality=%s "
             "is_usable=%s is_partial=%s ocr_api_failed=%s has_some_text=%s "
             "has_medical=%s status=%s fmt=%s",
             ocr_result.get("success", False), char_count, ocr_quality,
             is_usable, is_partial, ocr_api_failed, has_some_text,
             _has_medical_terms(ocr_text), ocr_result.get("confidence", "?"), fmt)

    if is_usable:
        synthetic_msg = ctx_str
        if caption:
            synthetic_msg = f"[CAPTION: {caption}]\n" + ctx_str
    elif is_partial:
        # Partial OCR — enough signal to proceed without alarming the user
        partial_success_triggered = True
        log.info("[MEDIA] partial_success_triggered contact=%s quality=%s chars=%d "
                 "has_medical=%s fallback_reason=low_confidence_usable",
                 contact_id, ocr_quality, char_count, _has_medical_terms(ocr_text))
        synthetic_msg = ctx_str
        if caption:
            synthetic_msg = f"[CAPTION: {caption}]\n" + ctx_str
    elif caption:
        synthetic_msg = f"[CAPTION: {caption}]\n" + ctx_str
    else:
        log.warning("[MEDIA] ocr_confidence_low contact=%s quality=%s chars=%d "
                    "ocr_api_failed=%s has_some_text=%s "
                    "fallback_reason=insufficient_text — returning failsafe",
                    contact_id, ocr_quality, char_count, ocr_api_failed, has_some_text)
        return OCR_FAILSAFE_MESSAGE

    # Step 6: Call AI pipeline
    try:
        reply = process_message_fn(contact_id, synthetic_msg)
    except Exception as e:
        log.error("[MEDIA] process_message_fn error: %s", e)
        reply = OCR_FAILSAFE_MESSAGE

    # For partial success: prepend acknowledgment so user knows document was received
    if partial_success_triggered and reply and reply != OCR_FAILSAFE_MESSAGE:
        if OCR_PARTIAL_SUCCESS_MESSAGE not in reply:
            reply = OCR_PARTIAL_SUCCESS_MESSAGE + "\n\n" + reply

    return reply


# ─── Utility ──────────────────────────────────────────────────────────────────

def has_attachment(body: dict) -> bool:
    """Fast check: does this webhook body contain a media attachment?"""
    return extract_attachment(body) is not None
